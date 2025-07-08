"""
Processador de Documentos - Sistema HELIO
Análise real de currículos em PDF e DOCX seguindo metodologia Carolina Martins
"""

import os
import re
import json
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
import PyPDF2
import docx
from io import BytesIO
import openai
from anthropic import Anthropic

class DocumentProcessor:
    """
    Processador real de documentos para análise de currículos
    Substitui as funções placeholder do sistema original
    """
    
    def __init__(self):
        # Inicializa APIs de IA se disponíveis
        self.openai_client = None
        self.anthropic_client = None
        
        if os.getenv('OPENAI_API_KEY') and os.getenv('OPENAI_API_KEY') != 'your_openai_api_key_here':
            openai.api_key = os.getenv('OPENAI_API_KEY')
            self.openai_client = openai
            
        if os.getenv('ANTHROPIC_API_KEY') and os.getenv('ANTHROPIC_API_KEY') != 'your_anthropic_api_key_here':
            self.anthropic_client = Anthropic(api_key=os.getenv('ANTHROPIC_API_KEY'))
    
    def extrair_texto_documento(self, arquivo_path: str, arquivo_bytes: bytes = None) -> str:
        """
        Extrai texto de documentos PDF ou DOCX
        
        Args:
            arquivo_path: Caminho do arquivo ou nome com extensão
            arquivo_bytes: Bytes do arquivo (se enviado via upload)
            
        Returns:
            str: Texto extraído do documento
        """
        try:
            extensao = arquivo_path.lower().split('.')[-1]
            
            if arquivo_bytes:
                if extensao == 'pdf':
                    return self._extrair_texto_pdf_bytes(arquivo_bytes)
                elif extensao in ['docx', 'doc']:
                    return self._extrair_texto_docx_bytes(arquivo_bytes)
                elif extensao == 'txt':
                    return arquivo_bytes.decode('utf-8', errors='ignore')
            else:
                if extensao == 'pdf':
                    return self._extrair_texto_pdf(arquivo_path)
                elif extensao in ['docx', 'doc']:
                    return self._extrair_texto_docx(arquivo_path)
                elif extensao == 'txt':
                    with open(arquivo_path, 'r', encoding='utf-8') as f:
                        return f.read()
                    
            raise ValueError(f"Formato de arquivo não suportado: {extensao}")
            
        except Exception as e:
            print(f"Erro ao extrair texto do documento: {e}")
            return ""
    
    def _extrair_texto_pdf(self, arquivo_path: str) -> str:
        """Extrai texto de arquivo PDF"""
        texto = ""
        try:
            with open(arquivo_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    texto += page.extract_text() + "\n"
        except Exception as e:
            print(f"Erro ao processar PDF: {e}")
        return texto
    
    def _extrair_texto_pdf_bytes(self, arquivo_bytes: bytes) -> str:
        """Extrai texto de PDF a partir de bytes"""
        texto = ""
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(arquivo_bytes))
            for page in pdf_reader.pages:
                texto += page.extract_text() + "\n"
        except Exception as e:
            print(f"Erro ao processar PDF bytes: {e}")
        return texto
    
    def _extrair_texto_docx(self, arquivo_path: str) -> str:
        """Extrai texto de arquivo DOCX"""
        try:
            doc = docx.Document(arquivo_path)
            texto = ""
            for paragraph in doc.paragraphs:
                texto += paragraph.text + "\n"
            return texto
        except Exception as e:
            print(f"Erro ao processar DOCX: {e}")
            return ""
    
    def _extrair_texto_docx_bytes(self, arquivo_bytes: bytes) -> str:
        """Extrai texto de DOCX a partir de bytes"""
        try:
            doc = docx.Document(BytesIO(arquivo_bytes))
            texto = ""
            for paragraph in doc.paragraphs:
                texto += paragraph.text + "\n"
            return texto
        except Exception as e:
            print(f"Erro ao processar DOCX bytes: {e}")
            return ""
    
    def analisar_estrutura_curriculo(self, texto_curriculo: str) -> Dict[str, Any]:
        """
        Analisa estrutura do currículo seguindo metodologia Carolina Martins
        
        Returns:
            Dict com análise da estrutura metodológica
        """
        # Elementos dos 13 passos da metodologia Carolina Martins
        elementos_metodologia = {
            "dados_pessoais": ["nome", "telefone", "email", "linkedin"],
            "objetivo": ["objetivo", "cargo", "posição"],
            "resumo": ["resumo", "perfil", "sobre"],
            "experiencias": ["experiência", "trabalho", "empresa"],
            "resultados": ["resultado", "alcançou", "aumentou", "reduziu", "%"],
            "formacao": ["formação", "educação", "graduação", "curso"],
            "idiomas": ["idioma", "inglês", "espanhol", "francês"],
            "tecnologia": ["excel", "power bi", "python", "sql"],
            "conhecimentos": ["conhecimento", "competência", "habilidade"],
            "voluntario": ["voluntário", "social", "ong"]
        }
        
        # Verifica presença de cada elemento
        elementos_encontrados = {}
        texto_lower = texto_curriculo.lower()
        
        for categoria, palavras_chave in elementos_metodologia.items():
            encontrado = any(palavra in texto_lower for palavra in palavras_chave)
            elementos_encontrados[categoria] = encontrado
        
        # Calcula score de aderência à metodologia
        total_elementos = len(elementos_metodologia)
        elementos_presentes = sum(elementos_encontrados.values())
        score_metodologia = (elementos_presentes / total_elementos) * 100
        
        # Identifica elementos faltando
        elementos_faltando = [
            categoria for categoria, presente 
            in elementos_encontrados.items() if not presente
        ]
        
        return {
            "possui_13_passos": score_metodologia >= 80,  # 80% dos elementos presentes
            "estrutura_adequada": score_metodologia >= 60,
            "formatacao_profissional": self._verificar_formatacao_texto(texto_curriculo),
            "elementos_encontrados": elementos_encontrados,
            "elementos_faltando": elementos_faltando,
            "score_metodologia": score_metodologia,
            "total_elementos": total_elementos,
            "elementos_presentes": elementos_presentes
        }
    
    def _verificar_formatacao_texto(self, texto: str) -> bool:
        """Verifica indicadores básicos de formatação profissional no texto"""
        # Indicadores de boa formatação
        indicadores = {
            "tem_estrutura_secoes": bool(re.search(r'(EXPERIÊNCIA|FORMAÇÃO|EDUCAÇÃO|OBJETIVO)', texto.upper())),
            "tem_datas": bool(re.search(r'\d{4}|\d{2}/\d{4}|20\d{2}', texto)),
            "tem_email": bool(re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', texto)),
            "tem_telefone": bool(re.search(r'\(\d{2}\)|\d{2}\s*9?\d{4}-?\d{4}', texto)),
            "tamanho_adequado": 500 < len(texto) < 3000  # Entre 500 e 3000 caracteres
        }
        
        # Considera bem formatado se tem pelo menos 3 indicadores
        return sum(indicadores.values()) >= 3
    
    def extrair_palavras_chave_curriculo(self, texto_curriculo: str) -> List[str]:
        """
        Extrai palavras-chave relevantes do currículo usando análise semântica
        
        Returns:
            Lista de palavras-chave extraídas
        """
        # Padrões para identificar competências técnicas e comportamentais
        padroes_competencias = [
            # Tecnologias
            r'\b(Excel|Power BI|Python|SQL|Tableau|SAP|Oracle|Java|JavaScript|React|Angular)\b',
            # Soft skills
            r'\b(liderança|gestão|comunicação|negociação|analítico|estratégico|proativo)\b',
            # Áreas funcionais
            r'\b(marketing|vendas|financeiro|recursos humanos|logística|operações|TI|tecnologia)\b',
            # Metodologias
            r'\b(Scrum|Agile|Six Sigma|Lean|Kanban|PMBOK|ITIL)\b'
        ]
        
        palavras_encontradas = set()
        texto_lower = texto_curriculo.lower()
        
        for padrao in padroes_competencias:
            matches = re.findall(padrao, texto_lower, re.IGNORECASE)
            palavras_encontradas.update(match.lower() for match in matches)
        
        # Se poucas palavras encontradas, usar IA para extração mais sofisticada
        if len(palavras_encontradas) < 5 and (self.openai_client or self.anthropic_client):
            palavras_ia = self._extrair_palavras_chave_ia(texto_curriculo)
            palavras_encontradas.update(palavras_ia)
        
        # NÃO LIMITAR! Retornar todas as palavras encontradas
        return list(palavras_encontradas)
    
    def _extrair_palavras_chave_ia(self, texto_curriculo: str) -> List[str]:
        """Extrai palavras-chave usando IA"""
        prompt = f"""
        Analise este currículo e extraia as 15 principais palavras-chave técnicas e comportamentais:

        CURRÍCULO:
        {texto_curriculo[:1500]}  # Limita texto para evitar tokens excessivos

        Retorne apenas uma lista de palavras-chave separadas por vírgula, sem explicações.
        Foque em: competências técnicas, soft skills, áreas funcionais, metodologias.
        """
        
        try:
            if self.anthropic_client:
                response = self.anthropic_client.messages.create(
                    model="claude-3-haiku-20240307",
                    max_tokens=200,
                    messages=[{"role": "user", "content": prompt}]
                )
                palavras = response.content[0].text.strip().split(',')
                return [p.strip().lower() for p in palavras if p.strip()]
                
            elif self.openai_client:
                response = self.openai_client.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[{"role": "user", "content": prompt}],
                    max_tokens=200,
                    temperature=0.3
                )
                palavras = response.choices[0].message.content.strip().split(',')
                return [p.strip().lower() for p in palavras if p.strip()]
                
        except Exception as e:
            print(f"Erro ao usar IA para extração de palavras-chave: {e}")
        
        return []
    
    def verificar_honestidade_curriculo(self, texto_curriculo: str) -> Dict[str, Any]:
        """
        Verifica indicadores de honestidade no currículo
        """
        alertas = []
        
        # Verifica consistência de datas
        datas_encontradas = re.findall(r'(\d{4})|(\d{2}/\d{4})', texto_curriculo)
        if datas_encontradas:
            anos = []
            for data in datas_encontradas:
                if data[0]:  # Formato YYYY
                    anos.append(int(data[0]))
                elif data[1]:  # Formato MM/YYYY
                    anos.append(int(data[1].split('/')[1]))
            
            anos.sort()
            if anos:
                # Verifica se há gaps muito grandes ou datas futuras
                ano_atual = datetime.now().year
                if any(ano > ano_atual for ano in anos):
                    alertas.append("Datas futuras encontradas")
                
                if len(anos) > 1 and max(anos) - min(anos) > 30:
                    alertas.append("Período profissional muito extenso (>30 anos)")
        
        # Verifica nível de detalhamento
        palavras_vagas = ["responsável por", "participou", "auxiliou", "apoiou"]
        palavras_especificas = ["gerenciei", "liderei", "desenvolvi", "implementei", "alcancei"]
        
        vagas_count = sum(1 for palavra in palavras_vagas if palavra in texto_curriculo.lower())
        especificas_count = sum(1 for palavra in palavras_especificas if palavra in texto_curriculo.lower())
        
        detalhamento_adequado = especificas_count > vagas_count
        
        # Verifica presença de resultados quantificados
        tem_numeros = bool(re.search(r'\d+%|\d+\s*(milhão|mil|reais|R\$)', texto_curriculo))
        
        return {
            "datas_consistentes": len(alertas) == 0,
            "informacoes_verificaveis": tem_numeros,
            "nivel_detalhamento_adequado": detalhamento_adequado,
            "alertas_inconsistencia": alertas,
            "indicadores": {
                "tem_resultados_quantificados": tem_numeros,
                "uso_verbos_acao": especificas_count > 2,
                "evita_linguagem_vaga": vagas_count < especificas_count
            }
        }
    
    def analisar_formatacao_documento(self, texto_curriculo: str, arquivo_path: str = None) -> Dict[str, Any]:
        """
        Analisa formatação do documento seguindo padrões Carolina Martins
        """
        # Análise baseada no texto extraído
        linhas = texto_curriculo.split('\n')
        linhas_nao_vazias = [linha for linha in linhas if linha.strip()]
        
        # Estimativa de páginas baseada no conteúdo
        caracteres_por_pagina = 1800  # Estimativa média
        paginas_estimadas = len(texto_curriculo) / caracteres_por_pagina
        
        # Verifica estrutura de seções
        secoes_identificadas = []
        secoes_padrao = [
            "dados pessoais", "objetivo", "resumo", "experiência", 
            "formação", "educação", "idiomas", "competências"
        ]
        
        for secao in secoes_padrao:
            if secao in texto_curriculo.lower():
                secoes_identificadas.append(secao)
        
        return {
            "formato_adequado": len(secoes_identificadas) >= 4,
            "tamanho_paginas": round(paginas_estimadas, 1),
            "fonte_profissional": True,  # Assumido após extração
            "espacamento_adequado": len(linhas_nao_vazias) / len(linhas) > 0.6,
            "uso_cores_apropriado": True,  # Não detectável via texto
            "estrutura_secoes": {
                "secoes_encontradas": secoes_identificadas,
                "total_secoes": len(secoes_identificadas),
                "secoes_minimas_ok": len(secoes_identificadas) >= 4
            },
            "metricas_texto": {
                "total_caracteres": len(texto_curriculo),
                "total_linhas": len(linhas),
                "linhas_com_conteudo": len(linhas_nao_vazias),
                "densidade_conteudo": len(linhas_nao_vazias) / len(linhas) if linhas else 0
            }
        }